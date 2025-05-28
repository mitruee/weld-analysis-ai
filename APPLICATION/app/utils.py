import asyncio
from fastapi import HTTPException
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import cv2
import numpy as np
from pathlib import Path


def ndarray_to_bytes(image_array: np.ndarray, format: str = "jpg") -> bytes:
    if format.lower() == "jpg":
        ext = ".jpg"
        encode_params = [int(cv2.IMWRITE_JPEG_QUALITY), 90]  # качество 90%
    elif format.lower() == "png":
        ext = ".png"
        encode_params = [int(cv2.IMWRITE_PNG_COMPRESSION), 6]  # компрессия 6
    else:
        raise ValueError("Unsupported image format. Use 'jpg' or 'png'")

    success, encoded_image = cv2.imencode(ext, image_array, encode_params)
    if not success:
        raise ValueError(f"Could not encode image to {format.upper()} format")

    return encoded_image.tobytes()

async def run_visualization_async(image_path: str, weights: str = "last.pt", conf: float = 0.15):
    try:
        process = await asyncio.create_subprocess_exec(
            "python",
            "visualize_predictions.py",
            image_path,
            "--weights", weights,
            "--conf", str(conf),
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )

        stdout, stderr = await process.communicate()

        if process.returncode != 0:
            raise HTTPException(
                status_code=500,
                detail=f"Ошибка выполнения: {stderr.decode()}"
            )

        return stdout.decode()
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Системная ошибка: {str(e)}"
        )

def _slice_panorama(img: np.ndarray) -> list[np.ndarray]:
    """Нарезка панорамы на тайлы"""
    SIZE_MAP = {
        (31920, 1152): 28,
        (30780, 1152): 27,
        (18144, 1142): 16,
    }
    h, w = img.shape[:2]
    tiles = SIZE_MAP.get((w, h))
    if tiles is None:
        raise ValueError(f"Неизвестный размер панорамы {w}×{h}")
    tw = w // tiles
    return [img[:, i * tw:(i + 1) * tw] for i in range(tiles)]


def create_defects_report(data, output_filename="static/reports/defects_report.docx"):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    title = doc.add_heading('Отчет о дефектах', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date_paragraph = doc.add_paragraph(f"Дата создания отчета: {current_time}")
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Тип дефекта'
    hdr_cells[2].text = 'Уверенность'
    hdr_cells[3].text = 'Индекс'
    hdr_cells[4].text = 'Координаты'
    hdr_cells[5].text = 'Длина по линейке'

    for cell in hdr_cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    defect_counter = 1
    for item in data:
        if item["status"] == "success" and item["defects"]:
            for defect in item["defects"]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(defect_counter)
                row_cells[1].text = defect["class"]
                row_cells[2].text = defect["confidence"]
                row_cells[3].text = str(defect["index"])
                row_cells[4].text = defect["coordinates"]
                row_cells[5].text = str(defect["length"])
                defect_counter += 1

    total_defects = defect_counter - 1
    stats_paragraph = doc.add_paragraph()
    stats_paragraph.add_run("Статистика:\n").bold = True
    stats_paragraph.add_run(f"Всего обнаружено дефектов: {total_defects}\n")

    defect_types = {}
    for item in data:
        if item["status"] == "success" and item["defects"]:
            for defect in item["defects"]:
                defect_type = defect["class"]
                defect_types[defect_type] = defect_types.get(defect_type, 0) + 1

    for defect_type, count in defect_types.items():
        stats_paragraph.add_run(f"{defect_type}: {count}\n")

    doc.add_paragraph("\n")
    sign_paragraph = doc.add_paragraph("Ответственный: _________________________")
    sign_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.save(output_filename)